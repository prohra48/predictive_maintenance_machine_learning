from __future__ import annotations

import os
from datetime import date
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from config import (
    CLEAN_DATA_FILE,
    EXPERIMENTS_DIR,
    FIGURES_DIR,
    LOGS_DIR,
    MODELS_DIR,
    REPORTS_DIR,
    REPORT_TITLE,
    TABLES_DIR,
    ensure_directories,
)
from utils import read_json


REPORT_DOCX = REPORTS_DIR / "Predictive_Maintenance_Final_Report.docx"
REPORT_PDF = REPORTS_DIR / "Predictive_Maintenance_Final_Report.pdf"


def money_quote(metric: float) -> str:
    return f"{metric:.2%}"


def load_inputs() -> dict:
    required_files = [
        LOGS_DIR / "data_registration_summary.json",
        LOGS_DIR / "eda_summary.json",
        LOGS_DIR / "data_preparation_summary.json",
        MODELS_DIR / "model_metadata.json",
        TABLES_DIR / "table_05_model_results.csv",
        TABLES_DIR / "table_07_classification_report.csv",
    ]
    missing = [str(path) for path in required_files if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "Run scripts 01 through 04 before generating the report. Missing: " + ", ".join(missing)
        )

    return {
        "registration": read_json(LOGS_DIR / "data_registration_summary.json"),
        "eda": read_json(LOGS_DIR / "eda_summary.json"),
        "prep": read_json(LOGS_DIR / "data_preparation_summary.json"),
        "model": read_json(MODELS_DIR / "model_metadata.json"),
        "results": pd.read_csv(TABLES_DIR / "table_05_model_results.csv"),
        "classification": pd.read_csv(TABLES_DIR / "table_07_classification_report.csv"),
        "overview": pd.read_csv(TABLES_DIR / "table_01_data_overview.csv"),
        "target": pd.read_csv(TABLES_DIR / "table_02_target_distribution.csv"),
        "feature_means": pd.read_csv(TABLES_DIR / "table_03_feature_means_by_condition.csv"),
    }


def hf_or_placeholder(env_name: str, label: str) -> str:
    value = os.getenv(env_name)
    if value:
        return value
    return f"To be filled after publishing {label}"


def core_report_text(data: dict) -> dict[str, list[str]]:
    model = data["model"]
    eda = data["eda"]
    prep = data["prep"]
    best = model["test_metrics"]
    target_rows = data["target"].to_dict(orient="records")
    normal = next((row for row in target_rows if int(row["engine_condition"]) == 0), {})
    maintenance = next((row for row in target_rows if int(row["engine_condition"]) == 1), {})

    top_importance = model.get("feature_importance", [])
    if top_importance:
        top_feature = top_importance[0]["feature"].replace("_", " ").title()
    else:
        top_feature = "the most informative engine sensor readings"

    return {
        "executive_summary": [
            "This project develops a machine learning solution that classifies whether an engine is operating normally or needs maintenance based on RPM, oil pressure, fuel pressure, coolant pressure, oil temperature, and coolant temperature.",
            f"The final model selected by test F1 score is {model['best_model']}. On the holdout test set, it achieved accuracy {money_quote(best['accuracy'])}, precision {money_quote(best['precision'])}, recall {money_quote(best['recall'])}, F1 score {money_quote(best['f1'])}, and ROC-AUC {money_quote(best['roc_auc'])}.",
            "The business value is earlier identification of likely engine issues, which can reduce unplanned breakdowns, improve maintenance scheduling, and give fleet/service teams a data-backed triage signal.",
        ],
        "business_context": [
            "Unexpected engine failure creates repair cost, operational downtime, safety risk, and poor customer experience. Predictive maintenance uses sensor readings to detect risk before failure happens.",
            "The objective is to convert historical engine sensor data into a classification model that can support proactive maintenance decisions.",
        ],
        "data_registration": [
            "A master project folder was created with a dedicated data subfolder. The raw CSV was stored under data/raw and the processed train/test datasets were stored under data/processed.",
            f"Raw data registration status: {data['registration']['hugging_face_dataset_status']}.",
            f"Hugging Face dataset link: {hf_or_placeholder('HF_DATASET_REPO', 'the dataset repo')}.",
        ],
        "eda": [
            f"The dataset contains {eda['rows']:,} records and {eda['columns']} columns. There are {eda['missing_values']} missing values and {eda['duplicates']} duplicate rows after cleaning.",
            f"The target variable is moderately imbalanced: {int(normal.get('records', 0)):,} normal records ({normal.get('percentage', 0):.2f}%) and {int(maintenance.get('records', 0)):,} maintenance records ({maintenance.get('percentage', 0):.2f}%).",
            "The EDA uses summary statistics, target distribution, univariate histograms, boxplots by condition, and a correlation heatmap to compare engine behavior under normal and maintenance states.",
        ],
        "data_preparation": [
            "Column names were standardized to snake_case, numeric fields were validated, duplicate rows were checked, and no unnecessary columns were retained. After Hugging Face publication, setting USE_HF_DATA=1 makes the scripts load the raw/train/test files directly from the Hugging Face dataset repo.",
            f"The data was split into {prep['train_rows']:,} training rows and {prep['test_rows']:,} testing rows using a stratified split so the target distribution remains consistent.",
            f"Processed data upload status: {prep['hugging_face_dataset_status']}.",
        ],
        "modeling": [
            "Five classification algorithms were evaluated: Decision Tree, Random Forest, Gradient Boosting, AdaBoost, and XGBoost. Hyperparameters were tuned with cross-validated grid search and the tuned parameters were logged in experiments/experiment_tracking.csv.",
            f"The best model is {model['best_model']}, selected primarily by F1 score because the business problem requires a balance between catching maintenance cases and avoiding excessive false alarms.",
            f"The strongest model signal is led by {top_feature}, based on the saved feature-importance chart. This helps translate the model output into practical engineering signals.",
            f"Model hub status: {model['hugging_face_model_status']}.",
        ],
        "deployment": [
            "A Streamlit app is provided in deployment/app.py. It accepts the six sensor inputs, loads the saved model, and returns the predicted engine condition with the estimated maintenance probability when available.",
            "deployment/requirements.txt lists the runtime dependencies. deployment/Dockerfile defines a containerized configuration for reproducible hosting.",
            f"Hugging Face Space link: {hf_or_placeholder('HF_SPACE_REPO', 'the Streamlit Space')}.",
        ],
        "github_actions": [
            ".github/workflows/pipeline.yml defines an automated workflow that installs dependencies, runs data registration, EDA, data preparation, model training, report generation, and artifact upload.",
            "If Hugging Face secrets are added in GitHub, the workflow can also upload datasets/models and push the Streamlit app files to a Hugging Face Space.",
            "Required GitHub secrets: HF_TOKEN, HF_DATASET_REPO, HF_MODEL_REPO, and HF_SPACE_REPO.",
        ],
        "output_evaluation": [
            "GitHub repository link: to be filled after pushing this project to GitHub.",
            "GitHub Actions screenshot: insert screenshot after a successful workflow run.",
            "Hugging Face Space screenshot: insert screenshot after the Streamlit app is published and tested.",
        ],
        "recommendations": [
            "Use the model as a triage signal rather than a full replacement for mechanical inspection.",
            "Prioritize recall for maintenance cases when the cost of missing an engine issue is higher than the cost of an extra inspection.",
            "Monitor model performance after deployment because operating ranges may shift across engine types, seasons, and usage patterns.",
            "Collect service-confirmed labels over time and retrain the model periodically to keep predictions aligned with real maintenance outcomes.",
        ],
    }


def add_docx_table(document: Document, df: pd.DataFrame, max_rows: int = 8) -> None:
    visible = df.head(max_rows).copy()
    table = document.add_table(rows=1, cols=len(visible.columns))
    table.style = "Table Grid"
    for i, column in enumerate(visible.columns):
        table.rows[0].cells[i].text = str(column)
    for _, row in visible.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            if isinstance(value, float):
                cells[i].text = f"{value:.4f}" if abs(value) <= 1 else f"{value:.2f}"
            else:
                cells[i].text = str(value)


def add_docx_figure(document: Document, figure_name: str, caption: str) -> None:
    path = FIGURES_DIR / figure_name
    if path.exists():
        document.add_picture(str(path), width=Inches(6.4))
        paragraph = document.add_paragraph(caption)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(data: dict, text: dict[str, list[str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_heading(REPORT_TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Final Project Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Generated on: {date.today().isoformat()}")

    sections = [
        ("Executive Summary", "executive_summary"),
        ("1. Business Context and Objective", "business_context"),
        ("2. Data Registration", "data_registration"),
        ("3. Exploratory Data Analysis", "eda"),
        ("4. Data Preparation", "data_preparation"),
        ("5. Model Building with Experimentation Tracking", "modeling"),
        ("6. Model Deployment", "deployment"),
        ("7. Automated GitHub Actions Workflow", "github_actions"),
        ("8. Output Evaluation", "output_evaluation"),
        ("9. Actionable Insights and Recommendations", "recommendations"),
    ]

    for heading, key in sections:
        document.add_heading(heading, level=1)
        if key in {"eda", "data_preparation", "modeling"}:
            document.add_heading("Methodology", level=2)
        for paragraph in text[key]:
            document.add_paragraph(paragraph)

        if key == "eda":
            document.add_heading("Table 1. Data Overview", level=2)
            add_docx_table(document, data["overview"].reset_index(drop=True), max_rows=7)
            document.add_heading("Table 2. Target Distribution", level=2)
            add_docx_table(document, data["target"], max_rows=5)
            add_docx_figure(document, "figure_01_target_distribution.png", "Figure 1. Engine Condition Distribution")
            add_docx_figure(document, "figure_02_feature_histograms.png", "Figure 2. Univariate Distribution of Engine Sensor Readings")
            add_docx_figure(document, "figure_03_boxplots_by_condition.png", "Figure 3. Sensor Readings by Engine Condition")
            add_docx_figure(document, "figure_04_correlation_heatmap.png", "Figure 4. Correlation Heatmap")

        if key == "modeling":
            document.add_heading("Table 3. Model Experiment Tracking Summary", level=2)
            add_docx_table(document, data["results"], max_rows=10)
            add_docx_figure(document, "figure_05_confusion_matrix.png", "Figure 5. Confusion Matrix for Best Model")
            add_docx_figure(document, "figure_06_feature_importance.png", "Figure 6. Feature Importance for Best Model")

    document.add_page_break()
    document.add_heading("Appendix A. Reproducible Local Commands", level=1)
    commands = [
        "pip install -r requirements.txt",
        "python src\\01_data_registration.py",
        "python src\\02_eda.py",
        "python src\\03_data_preparation.py",
        "python src\\04_model_training.py",
        "python src\\05_generate_report.py",
        "python src\\06_generate_notebook.py",
        "streamlit run deployment\\app.py",
    ]
    for command in commands:
        document.add_paragraph(command, style="List Bullet")

    document.add_heading("Appendix B. Code and Evidence Files", level=1)
    files = [
        "src/01_data_registration.py",
        "src/02_eda.py",
        "src/03_data_preparation.py",
        "src/04_model_training.py",
        "src/05_generate_report.py",
        "deployment/app.py",
        "deployment/Dockerfile",
        ".github/workflows/pipeline.yml",
        "experiments/experiment_tracking.csv",
        "models/model_metadata.json",
    ]
    for file_name in files:
        document.add_paragraph(file_name, style="List Bullet")

    document.save(REPORT_DOCX)


def as_pdf_table(df: pd.DataFrame, max_rows: int = 8) -> Table:
    visible = df.head(max_rows).copy()
    values = [list(visible.columns)]
    for _, row in visible.iterrows():
        formatted = []
        for value in row.tolist():
            if isinstance(value, float):
                formatted.append(f"{value:.4f}" if abs(value) <= 1 else f"{value:.2f}")
            else:
                formatted.append(str(value))
        values.append(formatted)
    table = Table(values, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F6F7")]),
            ]
        )
    )
    return table


def paragraph_list(items: Iterable[str], style: ParagraphStyle) -> list:
    flowables = []
    for item in items:
        flowables.append(Paragraph(item, style))
        flowables.append(Spacer(1, 0.08 * inch))
    return flowables


def add_pdf_figure(story: list, styles: dict, figure_name: str, caption: str) -> None:
    path = FIGURES_DIR / figure_name
    if path.exists():
        story.append(Image(str(path), width=6.5 * inch, height=3.7 * inch, kind="proportional"))
        story.append(Paragraph(caption, styles["Caption"]))
        story.append(Spacer(1, 0.12 * inch))


def page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(A4[0] - 0.55 * inch, 0.35 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(data: dict, text: dict[str, list[str]]) -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontSize=22,
            leading=28,
            alignment=TA_CENTER,
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#444444"),
            spaceAfter=8,
        )
    )

    doc = SimpleDocTemplate(
        str(REPORT_PDF),
        pagesize=A4,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.6 * inch,
    )

    story = [
        Spacer(1, 1.0 * inch),
        Paragraph(REPORT_TITLE, styles["CoverTitle"]),
        Paragraph("Final Project Report", styles["Title"]),
        Paragraph(f"Generated on: {date.today().isoformat()}", styles["Normal"]),
        PageBreak(),
    ]

    section_order = [
        ("Executive Summary", "executive_summary"),
        ("1. Business Context and Objective", "business_context"),
        ("2. Data Registration", "data_registration"),
        ("3. Exploratory Data Analysis", "eda"),
        ("4. Data Preparation", "data_preparation"),
        ("5. Model Building with Experimentation Tracking", "modeling"),
        ("6. Model Deployment", "deployment"),
        ("7. Automated GitHub Actions Workflow", "github_actions"),
        ("8. Output Evaluation", "output_evaluation"),
        ("9. Actionable Insights and Recommendations", "recommendations"),
    ]

    for heading, key in section_order:
        story.append(Paragraph(heading, styles["Heading1"]))
        if key in {"eda", "data_preparation", "modeling"}:
            story.append(Paragraph("Methodology", styles["Heading2"]))
        story.extend(paragraph_list(text[key], styles["BodyText"]))

        if key == "eda":
            story.append(Paragraph("Table 1. Data Overview", styles["Heading2"]))
            story.append(as_pdf_table(data["overview"].reset_index(drop=True), max_rows=7))
            story.append(Spacer(1, 0.12 * inch))
            story.append(Paragraph("Table 2. Target Distribution", styles["Heading2"]))
            story.append(as_pdf_table(data["target"], max_rows=5))
            story.append(Spacer(1, 0.12 * inch))
            add_pdf_figure(story, styles, "figure_01_target_distribution.png", "Figure 1. Engine Condition Distribution")
            add_pdf_figure(story, styles, "figure_02_feature_histograms.png", "Figure 2. Univariate Distribution of Engine Sensor Readings")
            add_pdf_figure(story, styles, "figure_03_boxplots_by_condition.png", "Figure 3. Sensor Readings by Engine Condition")
            add_pdf_figure(story, styles, "figure_04_correlation_heatmap.png", "Figure 4. Correlation Heatmap")

        if key == "modeling":
            story.append(Paragraph("Table 3. Model Experiment Tracking Summary", styles["Heading2"]))
            story.append(as_pdf_table(data["results"], max_rows=10))
            story.append(Spacer(1, 0.12 * inch))
            add_pdf_figure(story, styles, "figure_05_confusion_matrix.png", "Figure 5. Confusion Matrix for Best Model")
            add_pdf_figure(story, styles, "figure_06_feature_importance.png", "Figure 6. Feature Importance for Best Model")

    story.append(PageBreak())
    story.append(Paragraph("Appendix A. Reproducible Local Commands", styles["Heading1"]))
    story.extend(
        paragraph_list(
            [
                "pip install -r requirements.txt",
                "python src\\01_data_registration.py",
                "python src\\02_eda.py",
                "python src\\03_data_preparation.py",
                "python src\\04_model_training.py",
                "python src\\05_generate_report.py",
                "python src\\06_generate_notebook.py",
                "streamlit run deployment\\app.py",
            ],
            styles["Code"],
        )
    )

    story.append(Paragraph("Appendix B. Code and Evidence Files", styles["Heading1"]))
    story.extend(
        paragraph_list(
            [
                "src/01_data_registration.py",
                "src/02_eda.py",
                "src/03_data_preparation.py",
                "src/04_model_training.py",
                "src/05_generate_report.py",
                "deployment/app.py",
                "deployment/Dockerfile",
                ".github/workflows/pipeline.yml",
                "experiments/experiment_tracking.csv",
                "models/model_metadata.json",
            ],
            styles["BodyText"],
        )
    )

    doc.build(story, onFirstPage=page_number, onLaterPages=page_number)


def main() -> None:
    ensure_directories()
    data = load_inputs()
    text = core_report_text(data)
    build_docx(data, text)
    build_pdf(data, text)
    print(f"Report written to {REPORT_DOCX}")
    print(f"PDF written to {REPORT_PDF}")


if __name__ == "__main__":
    main()
